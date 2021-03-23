# -*- coding: utf-8 -*-
"""
使用python對Excel進行解析
"""
import numpy as np
import sys
from openpyxl import load_workbook
import os
from win32com import client
# 獲取桌面的路徑
def GetDesktopPath():
    return os.path.join(os.path.expanduser("~"), 'Desktop')

path = GetDesktopPath()  # 形成資料夾的路徑便後續重複使用
docname = input('輸入檔名')
output_doc = input('輸出檔名')
workbook = load_workbook(filename=path + "/%s" %docname)
sheet = workbook.active # 獲取當前頁
# 可以用程式碼獲取資料範圍，如果要批處理迴圈迭代也方便
# 獲取有資料範圍
print(sheet.dimensions)
maxR = len(sheet['A'])
total = maxR-1#總人數
### A1:Q36
###test
# cells = sheet['A1:A2']  # 返回A1-A4的4個單元格
# cells = sheet['C'] # 獲取A列
# cells = sheet['A:C'] # 獲取A-C列
# cells = sheet[5] # 獲取第5行
# print(cells)
# 注意如果是上述用cells獲取返回的是巢狀元祖
# for cell in cells:
#     print(cell[0].value) # 遍歷cells依然需要取出元祖中元素才可以獲取值

# 獲取一個範圍的所有cell
# 也可以用iter_col返回列
# for row in sheet.iter_rows(min_row=1, max_row=8,min_col=2, max_col=4):
#     for cell in row:
#         print(cell.value)
#性別(完成)
SUM_A_male=0
SUM_A_female=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=2, max_col=2):
    for cell in row:
        if cell.value=="男":
            SUM_A_male+=1
        # print(cell.value)
SUM_A_female=total-SUM_A_male
SUM_A_male_rate=(SUM_A_male/total)*100
SUM_A_female_rate=100-SUM_A_male_rate
print('%.1f'%SUM_A_male)
print('%.1f'%SUM_A_female)
print('%.1f'%SUM_A_male_rate)
print('%.1f'%SUM_A_female_rate)
print(total)
print("100%")
##b.學院(完成)
b = np.zeros(6)
b_rate = np.zeros(6)
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=3, max_col=3):
    for cell in row:
        if cell.value=="工學院":
            b[0]+=1
        elif cell.value=="理學院":
            b[1]+=1
        elif cell.value=="商管學院":
            b[2]+=1
        elif cell.value=="文學院":
            b[3]+=1
        elif cell.value=="外國語言學院":
            b[4]+=1
        elif cell.value=="教育學院":
            b[5]+=1
        # print(cell.value)
b_rate = b / total * 100
float_formatter = lambda x: "%.0f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})
print(b)
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})
print(b_rate)
# ###c.年級(完成)
c=np.zeros(5)
c_r = np.zeros(5)

for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=4, max_col=4):
    for cell in row:
        if cell.value=="一年級":
            c[0]+=1
        elif cell.value=="二年級":
            c[1]+=1
        elif cell.value=="三年級":
            c[2]+=1
        elif cell.value=="四年級":
            c[3]+=1
        else :
            c[4]+=1
c_r=c/total*100
        # print(cell.value)
float_formatter = lambda x: "%.0f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})
print(c)
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})
print(c_r)

# #3d
# d=np.zeros(2)
# d_r=np.zeros(2)
# SUM_d=0
# for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=5, max_col=5):
#     for cell in row:
#         if cell.value=="是":
#             d[0]+=1
#         elif cell.value=="否":
#             d[1]+=1

        
        
#         # print(cell.value)
# float_formatter = lambda x: "%d" % x
# np.set_printoptions(formatter={'float_kind':float_formatter})  
# print(d)
# d_r = d / total * 100
# float_formatter = lambda x: "%.1f" % x
# np.set_printoptions(formatter={'float_kind':float_formatter})      
# print(d_r)  
# SUM_d=SUM_d/total
# print("%.1f" %(SUM_d))
##課程內容(完成)
e=np.zeros(6)
e_r=np.zeros(6)
SUM_E=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=5, max_col=5):
    for cell in row:
        if cell.value==1:
            e[0]+=1
        elif cell.value==2:
            e[1]+=1
        elif cell.value==3:
            e[2]+=1
        elif cell.value==4:
            e[3]+=1
        elif cell.value==5:
            e[4]+=1
        elif cell.value==6:
            e[5]+=1
        SUM_E= SUM_E + cell.value
        
        
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(e)
e_r = e / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(e_r)  
SUM_E=SUM_E/total
print("%.1f" %(SUM_E))

##課師表現(完成)
f=np.zeros(6)
f_r=np.zeros(6)
SUM_f=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=6, max_col=6):
    for cell in row:
        if cell.value==1:
            f[0]+=1
        elif cell.value==2:
            f[1]+=1
        elif cell.value==3:
            f[2]+=1
        elif cell.value==4:
            f[3]+=1
        elif cell.value==5:
            f[4]+=1
        elif cell.value==6:
            f[5]+=1
        SUM_f+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(f)
f_r = f / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(f_r)  
SUM_f=SUM_f/total
print("%.1f" %(SUM_f))
##講師授課態度(完成)
g=np.zeros(6)
g_r=np.zeros(6)
SUM_g=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=7, max_col=7):
    for cell in row:
        if cell.value==1:
            g[0]+=1
        elif cell.value==2:
            g[1]+=1
        elif cell.value==3:
            g[2]+=1
        elif cell.value==4:
            g[3]+=1
        elif cell.value==5:
            g[4]+=1
        elif cell.value==6:
            g[5]+=1
        SUM_g+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(g)
g_r = g / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(g_r)  
SUM_g=SUM_g/total
print("%.1f" %(SUM_g))
##課程難易度(完成)
h=np.zeros(6)
h_r=np.zeros(6)
SUM_h=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=8, max_col=8):
    for cell in row:
        if cell.value==1:
            h[0]+=1
        elif cell.value==2:
            h[1]+=1
        elif cell.value==3:
            h[2]+=1
        elif cell.value==4:
            h[3]+=1
        elif cell.value==5:
            h[4]+=1
        elif cell.value==6:
            h[5]+=1
        SUM_h+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(h)
h_r = h / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(h_r)  
SUM_h=SUM_h/total
print("%.1f" %(SUM_h))

##器材數量(完成)
i=np.zeros(6)
i_r=np.zeros(6)
SUM_i=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=9, max_col=9):
    for cell in row:
        if cell.value==1:
            i[0]+=1
        elif cell.value==2:
            i[1]+=1
        elif cell.value==3:
            i[2]+=1
        elif cell.value==4:
            i[3]+=1
        elif cell.value==5:
            i[4]+=1
        elif cell.value==6:
            i[5]+=1
        SUM_i+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(i)
i_r = i / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(i_r)  
SUM_i=SUM_i/total
print("%.1f" %(SUM_i))

##教學設備(完成)
j=np.zeros(6)
j_r=np.zeros(6)
SUM_j=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=10, max_col=10):
    for cell in row:
        if cell.value==1:
            j[0]+=1
        elif cell.value==2:
            j[1]+=1
        elif cell.value==3:
            j[2]+=1
        elif cell.value==4:
            j[3]+=1
        elif cell.value==5:
            j[4]+=1
        elif cell.value==6:
            j[5]+=1
        SUM_j+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(j)
j_r = j / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(j_r)  
SUM_j=SUM_j/total
print("%.1f" %(SUM_j))
##活動整體
k=np.zeros(6)
k_r=np.zeros(6)
SUM_k=0
for row in sheet.iter_rows(min_row=2,max_row=maxR,min_col=11, max_col=11):
    for cell in row:
        if cell.value==1:
            k[0]+=1
        elif cell.value==2:
            k[1]+=1
        elif cell.value==3:
            k[2]+=1
        elif cell.value==4:
            k[3]+=1
        elif cell.value==5:
            k[4]+=1
        elif cell.value==6:
            k[5]+=1
        SUM_k+=cell.value
        # print(cell.value)
float_formatter = lambda x: "%d" % x
np.set_printoptions(formatter={'float_kind':float_formatter})  
print(k)
k_r = k / total * 100
float_formatter = lambda x: "%.1f" % x
np.set_printoptions(formatter={'float_kind':float_formatter})      
print(k_r)  
SUM_k=SUM_k/total
print("%.1f" %(SUM_k))

from docx.enum.text import WD_ALIGN_PARAGRAPH ##處理WORD置中
from docx.shared import Cm, Pt  #加入可調整的 word 單位
from docx.oxml.ns import qn #加入可調整的 word 單位
docx_path = path + '/' + output_doc
# doc轉docx的函式
def doc2docx(doc_path,docx_path):
    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(docx_path, 16)
    doc.Close()
    word.Quit()
    print('\n doc檔案已轉換為docx \n')
if not os.path.exists(docx_path):
    doc2docx(docx_path[:-1], docx_path)
docx_path = path + '/' + output_doc
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 例項化
document = Document(docx_path)
# 讀取word中的所有表格
tables = document.tables
# print(len(tables))
# 10
##table1
tables[0].cell(2, 1).text = str(SUM_A_male)
tables[0].cell(2, 2).text = str(SUM_A_female)
tables[0].cell(2, 3).text = str(total)
SUM_A_male_rate = str(SUM_A_male_rate)
SUM_A_female_rate = str(SUM_A_female_rate)
tables[0].cell(3, 1).text = SUM_A_male_rate[0:5]+"%"
tables[0].cell(3, 2).text = SUM_A_female_rate[0:5]+"%"
tables[0].cell(3, 3).text = "100%"
#table2
for column in range(1, 7):
    var_b = str(b[column-1])
    tables[1].cell(2, column).text = var_b[0:2]
    
tables[1].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_b_r = str(b_rate[column-1])
    tables[1].cell(3, column).text = var_b_r[0:4]+"%"
tables[1].cell(3, 7).text = "100%"
##table3
# for column in range(1, 3):
#     var_d = str(d[column-1])
#     tables[2].cell(2, column).text = var_d[0:2]
    
# tables[2].cell(2, 3).text = str(total)
# for column in range(1, 3):
#     var_d_r = str(d_r[column-1])
#     tables[2].cell(3, column).text = var_d_r[0:4]+"%"
# tables[2].cell(3, 3).text = "100%"
###
for column in range(1, 4):
    var_c = str(c[column-1])
    tables[3].cell(2, column).text = var_c[0:2]
    
tables[3].cell(2, 4).text = str(total)
for column in range(1, 4):
    var_c_r = str(c_r[column-1])
    tables[3].cell(3, column).text = var_c_r[0:4]+"%"
tables[3].cell(3, 4).text = "100%"
# SUM_d = str(SUM_d)

##table4
for column in range(1, 7):
    var_e = str(e[column-1])
    tables[4].cell(2, column).text = var_e[0:2]
 
###
tables[4].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_e_r = str(e_r[column-1])
    tables[4].cell(3, column).text = var_e_r[0:4]+"%"
tables[4].cell(3, 7).text = "100%"
SUM_E = str(SUM_E)
tables[4].cell(4, 1).text = SUM_E[0:3]
###table5
for column in range(1, 7):
    var_f = str(f[column-1])
    tables[5].cell(2, column).text = var_f[0:2]
    
tables[5].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_f_r = str(f_r[column-1])
    tables[5].cell(3, column).text = var_f_r[0:4]+"%"
tables[5].cell(3, 7).text = "100%"
SUM_f = str(SUM_f)
tables[5].cell(4, 1).text = SUM_f[0:3]

###table6
for column in range(1, 7):
    var_g = str(g[column-1])
    tables[6].cell(2, column).text = var_g[0:2]
    
tables[6].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_g_r = str(g_r[column-1])
    tables[6].cell(3, column).text = var_g_r[0:4]+"%"
tables[6].cell(3, 7).text = "100%"
SUM_g = str(SUM_g)
tables[6].cell(4, 1).text = SUM_g[0:3]

###table7
for column in range(1, 7):
    var_h = str(h[column-1])
    tables[7].cell(2, column).text = var_h[0:2]
    
tables[7].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_h_r = str(h_r[column-1])
    tables[7].cell(3, column).text = var_h_r[0:4]+"%"
tables[7].cell(3, 7).text = "100%"
SUM_h = str(SUM_h)
tables[7].cell(4, 1).text = SUM_h[0:3]

###table8
for column in range(1, 7):
    var_i = str(i[column-1])
    tables[8].cell(2, column).text = var_i[0:2]
    
tables[8].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_i_r = str(i_r[column-1])
    tables[8].cell(3, column).text = var_i_r[0:4]+"%"
tables[8].cell(3, 7).text = "100%"
SUM_i = str(SUM_i)
tables[8].cell(4, 1).text = SUM_i[0:3]

###table9
for column in range(1, 7):
    var_j = str(j[column-1])
    tables[9].cell(2, column).text = var_j[0:2]
    
tables[9].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_j_r = str(j_r[column-1])
    tables[9].cell(3, column).text = var_j_r[0:4]+"%"
tables[9].cell(3, 7).text = "100%"
SUM_j = str(SUM_j)
tables[9].cell(4, 1).text = SUM_j[0:3]

###table10
for column in range(1, 7):
    var_k = str(k[column-1])
    tables[10].cell(2, column).text = var_k[0:2]
    
tables[10].cell(2, 7).text = str(total)
for column in range(1, 7):
    var_k_r = str(k_r[column-1])
    tables[10].cell(3, column).text = var_k_r[0:4]+"%"
tables[10].cell(3, 7).text = "100%"
SUM_k = str(SUM_k)
tables[10].cell(4, 1).text = SUM_k[0:3]

# for paragraph in document.paragraphs:
# for paragraph in document.paragraphs:
#     if paragraph.style.name.startswith('Normal Table'):
#         paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER # LEFT, CENTER, RIGHT
#         for run in paragraph.runs:
#             run.font.name = '標楷體'
#             run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
#             run.font.size = Pt(18)
#             print(run.text)

# def iter_items(paragraphs):
#     for paragraph in document.paragraphs:
#         if paragraph.style.name.startswith('Agt'):
#             yield paragraph
#         if paragraph.style.name.startswith('TOC'):
#             yield paragraph
#         if paragraph.style.name.startswith('Heading'):
#             yield paragraph
#         if paragraph.style.name.startswith('Title'):
#             yield paragraph
#         if paragraph.style.name.startswith('Heading'):
#             yield paragraph
#         if paragraph.style.name.startswith('Table Normal'):
#             yield paragraph
#         if paragraph.style.name.startswith('List'):
#             yield paragraph
    
# for item in iter_items(document.paragraphs):
#     print (item.style)
# for style in document.styles:
#     print(style.name)

document.save(path + '/'  + output_doc)
print('\n檔案已生成')